# -*- coding: utf-8 -*-
"""Gera a versao editavel (.docx) da apresentacao do MVP."""

import os

import pandas as pd
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DIR_BASE = os.path.dirname(os.path.abspath(__file__))
DIR_SAIDA = os.path.join(DIR_BASE, "output")
CAMINHO_DOCX = os.path.join(DIR_BASE, "apresentacao_sompo_risco_python.docx")

AZUL = RGBColor(0x1F, 0x3A, 0x5F)


def set_cell_shading(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def tabela(doc, cabecalho, linhas, larguras=None):
    t = doc.add_table(rows=1, cols=len(cabecalho))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, texto in enumerate(cabecalho):
        celula = t.rows[0].cells[i]
        celula.text = ""
        p = celula.paragraphs[0]
        r = p.add_run(texto)
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.size = Pt(9.5)
        set_cell_shading(celula, "1F3A5F")
    for linha in linhas:
        celulas = t.add_row().cells
        for i, valor in enumerate(linha):
            celulas[i].text = ""
            p = celulas[i].paragraphs[0]
            r = p.add_run(str(valor))
            r.font.size = Pt(9.5)
    if larguras:
        for row in t.rows:
            for i, w in enumerate(larguras):
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return t


def montar():
    df = pd.read_csv(os.path.join(DIR_SAIDA, "resultados.csv"), encoding="utf-8")

    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10.5)

    titulo = doc.add_heading("Sompo Field Risk", level=0)
    for run in titulo.runs:
        run.font.color.rgb = AZUL

    sub = doc.add_paragraph("MVP de Análise de Risco Operacional (Python)")
    sub.runs[0].italic = True

    doc.add_paragraph("Disciplina: Computational Thinking with Python - Sprint 3")
    doc.add_paragraph("Professor: Kévin Allan Sales Rodrigues")
    doc.add_paragraph("Projeto: Sompo Seguros")

    doc.add_heading("Objetivo", level=1)
    doc.add_paragraph(
        "MVP em Python que recebe dados operacionais de equipamentos (telemetria real ou "
        "simulada), processa essas informações num motor de risco e gera saídas "
        "interpretáveis: score, classificação e alertas. É a primeira versão do backend de "
        "análise de risco do projeto Sompo Seguros nessa disciplina, conectando a entrada "
        "de dados ao modelo de risco."
    )

    doc.add_heading("Estrutura do projeto", level=1)
    tabela(doc, ["Arquivo", "Responsabilidade"], [
        ["main.py", "Pipeline: entrada -> processamento -> saída"],
        ["src/entrada.py", "Leitura da planilha e validação de dados obrigatórios"],
        ["src/sensores.py", "Simulação de sensores/API (telemetria, ambiente, operação)"],
        ["src/risco.py", "Cálculo do score, classificação, fator principal e alertas"],
        ["src/saida.py", "Resumo no console, exportações (CSV/JSON) e gráfico"],
    ], larguras=[5, 11])

    doc.add_page_break()

    doc.add_heading("Regras de negócio implementadas", level=1)
    doc.add_paragraph(
        "O score de risco (0-100) é calculado a partir de quatro indicadores de "
        "telemetria, com pesos definidos para este MVP:"
    )
    tabela(doc, ["Indicador", "Direção", "Peso máximo"], [
        ["Estilo de condução (nota 0-10)", "invertido", "40 pontos"],
        ["Estilo de condução na frenagem (nota 0-10)", "invertido", "30 pontos"],
        ["Grau de dificuldade da rota (nota 0-10)", "direto", "15 pontos"],
        ["Desaceleração / total percorrido (%)", "direto", "~15 pontos"],
    ], larguras=[9, 3.5, 3.5])
    tabela(doc, ["Faixa de score", "Classificação"], [
        ["0 - 25", "BAIXO"],
        ["26 - 50", "MODERADO"],
        ["51 - 75", "ALTO"],
        ["76 - 100", "CRITICO"],
    ], larguras=[8, 8])

    doc.add_page_break()

    doc.add_heading("Resultados - rodada com dados reais (FleetBoard)", level=1)
    total = len(df)
    contagem = df["classificacao"].value_counts()
    doc.add_paragraph(f"Total de equipamentos processados: {total}")
    tabela(doc, ["Classificação", "Quantidade", "% da frota"], [
        [rotulo, contagem.get(rotulo, 0), f"{contagem.get(rotulo, 0) / total * 100:.1f}%"]
        for rotulo in ["BAIXO", "MODERADO", "ALTO", "CRITICO"]
    ], larguras=[6, 5, 5])

    caminho_grafico = os.path.join(DIR_SAIDA, "distribuicao_risco.png")
    if os.path.exists(caminho_grafico):
        doc.add_picture(caminho_grafico, width=Cm(13))

    doc.add_heading("Principais fatores de risco na frota", level=1)
    ranking = df["fator_principal"].value_counts()
    tabela(
        doc, ["Fator", "Equipamentos onde foi o principal fator"],
        [[fator, qtd] for fator, qtd in ranking.items()],
        larguras=[9, 7],
    )

    doc.add_heading("Equipamentos com maior risco identificado", level=1)
    top6 = df.sort_values("score_risco", ascending=False).head(6)
    tabela(
        doc, ["Equipamento", "Período", "Score", "Classificação", "Fator principal"],
        [
            [l["equipamento_id"], l["periodo"], l["score_risco"], l["classificacao"], l["fator_principal"]]
            for _, l in top6.iterrows()
        ],
        larguras=[4, 2.7, 1.8, 3, 4.5],
    )

    doc.add_page_break()

    doc.add_heading("Requisitos técnicos atendidos", level=1)
    tabela(doc, ["Requisito", "Onde está"], [
        ["Funções para modularizar o fluxo", "entrada.py, sensores.py, risco.py, saida.py"],
        ["Estruturas condicionais (validação, classificação, alertas)", "validar_dados, classificar_risco, gerar_alerta"],
        ["Uso de pandas", "Leitura, limpeza e agregação dos dados"],
        ["Pipeline claro (entrada / processamento / saída)", "main.py"],
        ["README detalhado no GitHub", "Ver link do repositório abaixo"],
    ], larguras=[9, 7])

    doc.add_heading("Cobertura funcional do enunciado", level=1)
    tabela(doc, ["Item pedido", "Status"], [
        ["Backend modular com funções", "OK"],
        ["Integração com o modelo de risco", "OK"],
        ["Entrada e validação de dados (reais ou simulados)", "OK"],
        ["Simulação de sensores/API (telemetria, ambiente, operação)", "OK"],
        ["Saídas interpretáveis (score, classificação, alertas)", "OK"],
        ["Relatórios/dashboards simples e fatores de risco", "OK"],
        ["Consulta rápida dos resultados", "OK"],
        ["Validação funcional do MVP (testado end-to-end)", "OK"],
    ], larguras=[11, 5])

    doc.add_heading("Código-fonte", level=1)
    doc.add_paragraph(
        "Repositório (privado, acesso mediante convite): "
        "github.com/susuelen153-design/sompo-risco-python"
    )

    doc.save(CAMINHO_DOCX)
    print("DOCX gerado em:", CAMINHO_DOCX)


if __name__ == "__main__":
    montar()
